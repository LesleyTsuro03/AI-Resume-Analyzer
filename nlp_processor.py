import spacy
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
import re
from PyPDF2 import PdfReader
from docx import Document
import pandas as pd
from typing import Dict, List, Tuple
import string
from dateutil import parser
import phonenumbers
from email_validator import validate_email, EmailNotValidError
import hashlib
import uuid
from datetime import datetime

# Download required NLTK data
nltk.download('punkt', quiet=True)
nltk.download('stopwords', quiet=True)
nltk.download('averaged_perceptron_tagger', quiet=True)
nltk.download('maxent_ne_chunker', quiet=True)
nltk.download('words', quiet=True)

class AdvancedResumeAnalyzer:
    def __init__(self):
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            raise Exception("Please download the spaCy model: python -m spacy download en_core_web_sm")
        
        self.stop_words = set(stopwords.words('english'))
        
        # Comprehensive skills database covering ALL industries
        self.skill_keywords = {
            # TECHNICAL & IT SKILLS
            'programming_languages': [
                'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'go', 'rust', 'swift', 'kotlin',
                'php', 'ruby', 'scala', 'r', 'matlab', 'perl', 'html', 'css', 'sass', 'less', 'sql', 'pl/sql',
                'bash', 'shell', 'powershell', 'dart', 'assembly', 'fortran', 'cobol', 'visual basic', 'vba'
            ],
            'web_development': [
                'django', 'flask', 'fastapi', 'spring', 'express', 'react', 'angular', 'vue', 'svelte',
                'laravel', 'ruby on rails', 'asp.net', 'node.js', 'next.js', 'nuxt.js', 'jquery', 'bootstrap',
                'tailwind', 'webpack', 'babel', 'npm', 'yarn', 'graphql', 'rest api', 'soap', 'json', 'xml',
                'ajax', 'web services', 'microservices', 'api development'
            ],
            'mobile_development': [
                'android', 'ios', 'react native', 'flutter', 'xamarin', 'ionic', 'swiftui', 'kotlin multiplatform',
                'mobile app development', 'cross-platform development', 'android studio', 'xcode'
            ],
            'databases': [
                'mysql', 'postgresql', 'mongodb', 'redis', 'sqlite', 'oracle', 'sql server', 'cassandra',
                'dynamodb', 'elasticsearch', 'firebase', 'cosmosdb', 'mariadb', 'db2', 'sqlite', 'hbase',
                'neo4j', 'arangodb', 'couchbase', 'rethinkdb', 'database design', 'database management'
            ],
            'cloud_platforms': [
                'aws', 'azure', 'google cloud', 'gcp', 'docker', 'kubernetes', 'terraform', 'ansible',
                'jenkins', 'gitlab', 'github actions', 'heroku', 'digital ocean', 'linode', 'vultr',
                'ibm cloud', 'oracle cloud', 'alibaba cloud', 'openshift', 'cloudformation', 'cloud computing'
            ],
            'data_science_ai': [
                'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'keras', 'scikit-learn',
                'pandas', 'numpy', 'matplotlib', 'seaborn', 'plotly', 'tableau', 'power bi', 'qlik',
                'data analysis', 'data visualization', 'statistical analysis', 'natural language processing',
                'computer vision', 'neural networks', 'reinforcement learning', 'big data', 'hadoop', 'spark',
                'data mining', 'predictive modeling', 'artificial intelligence', 'nlp', 'computer vision',
                'data engineering', 'etl', 'data warehousing', 'business intelligence'
            ],
            'cybersecurity': [
                'network security', 'information security', 'cyber security', 'penetration testing', 'ethical hacking',
                'vulnerability assessment', 'siem', 'soc', 'firewall', 'vpn', 'encryption', 'cryptography',
                'incident response', 'threat intelligence', 'risk assessment', 'compliance', 'gdpr', 'hipaa',
                'pci dss', 'iso 27001', 'nist', 'owasp', 'security audit', 'digital forensics'
            ],
            'devops': [
                'ci/cd', 'continuous integration', 'continuous deployment', 'jenkins', 'gitlab ci', 'github actions',
                'ansible', 'puppet', 'chef', 'docker', 'kubernetes', 'helm', 'terraform', 'infrastructure as code',
                'monitoring', 'logging', 'prometheus', 'grafana', 'elk stack', 'splunk'
            ],
            
            # BUSINESS & MANAGEMENT SKILLS
            'project_management': [
                'project management', 'agile', 'scrum', 'kanban', 'waterfall', 'prince2', 'pmp', 'pmbok',
                'jira', 'trello', 'asana', 'basecamp', 'risk management', 'stakeholder management',
                'budget management', 'resource allocation', 'project planning', 'gantt chart', 'critical path',
                'ms project', 'smartsheet'
            ],
            'business_analysis': [
                'business analysis', 'requirements gathering', 'user stories', 'use cases', 'process modeling',
                'bpmn', 'uml', 'swot analysis', 'gap analysis', 'cost-benefit analysis', 'roi analysis',
                'kpi tracking', 'metrics', 'dashboard creation', 'business intelligence', 'process improvement'
            ],
            'finance_accounting': [
                'financial analysis', 'accounting', 'bookkeeping', 'financial modeling', 'budgeting', 'forecasting',
                'quickbooks', 'xero', 'sage', 'ifrs', 'gaap', 'tax preparation', 'audit', 'internal controls',
                'financial reporting', 'cash flow management', 'investment analysis', 'risk management',
                'sap fico', 'oracle financials', 'financial planning'
            ],
            'marketing_sales': [
                'digital marketing', 'social media marketing', 'seo', 'sem', 'google analytics', 'google ads',
                'facebook ads', 'content marketing', 'email marketing', 'marketing automation', 'hubspot',
                'market research', 'brand management', 'sales', 'business development', 'lead generation',
                'customer relationship management', 'crm', 'salesforce', 'negotiation', 'presentation skills',
                'copywriting', 'brand strategy', 'market analysis'
            ],
            'human_resources': [
                'recruitment', 'talent acquisition', 'onboarding', 'training and development', 'performance management',
                'compensation and benefits', 'employee relations', 'hr policies', 'labor law', 'succession planning',
                'organizational development', 'change management', 'diversity and inclusion', 'hr analytics',
                'payroll management', 'hris', 'workday', 'successfactors'
            ],
            'supply_chain': [
                'supply chain management', 'logistics', 'inventory management', 'procurement', 'purchasing',
                'warehouse management', 'demand planning', 'supplier management', 'sap mm', 'oracle scm'
            ],
            
            # CREATIVE & DESIGN SKILLS
            'graphic_design': [
                'adobe photoshop', 'adobe illustrator', 'adobe indesign', 'adobe xd', 'figma', 'sketch',
                'coreldraw', 'canva', 'ui design', 'ux design', 'user experience', 'user interface',
                'wireframing', 'prototyping', 'visual design', 'brand identity', 'logo design', 'typography',
                'color theory', 'print design', 'digital design', 'adobe creative suite'
            ],
            'video_audio': [
                'video editing', 'adobe premiere pro', 'final cut pro', 'after effects', 'davinci resolve',
                'motion graphics', 'animation', '3d modeling', 'blender', 'maya', 'cinema 4d', 'audio editing',
                'ableton live', 'logic pro', 'pro tools', 'sound design', 'podcast production'
            ],
            
            # PROFESSIONAL & SOFT SKILLS
            'soft_skills': [
                'communication', 'leadership', 'teamwork', 'problem solving', 'critical thinking',
                'adaptability', 'time management', 'creativity', 'collaboration', 'analytical skills',
                'strategic planning', 'negotiation', 'presentation', 'conflict resolution', 'decision making',
                'emotional intelligence', 'public speaking', 'writing', 'research', 'attention to detail',
                'interpersonal skills', 'customer service', 'mentoring', 'coaching'
            ],
            'languages': [
                'english', 'french', 'spanish', 'german', 'chinese', 'japanese', 'arabic', 'portuguese',
                'russian', 'hindi', 'shona', 'ndebele', 'swahili', 'zulu', 'afrikaans'
            ],
            
            # INDUSTRY-SPECIFIC SKILLS
            'healthcare': [
                'patient care', 'medical terminology', 'healthcare management', 'clinical research',
                'pharmaceutical', 'nursing', 'medical coding', 'hipaa compliance', 'electronic health records',
                'medical devices', 'healthcare analytics'
            ],
            'education': [
                'teaching', 'curriculum development', 'lesson planning', 'classroom management',
                'educational technology', 'student assessment', 'academic advising', 'research methodology'
            ],
            'legal': [
                'legal research', 'contract law', 'corporate law', 'litigation', 'legal writing',
                'compliance', 'intellectual property', 'legal documentation', 'case management'
            ],
            'engineering': [
                'mechanical engineering', 'electrical engineering', 'civil engineering', 'chemical engineering',
                'cad', 'autocad', 'solidworks', 'matlab', 'simulation', 'project engineering', 'quality control'
            ],
            'manufacturing': [
                'lean manufacturing', 'six sigma', 'quality assurance', 'production planning', 'supply chain',
                'process improvement', 'manufacturing operations', 'health and safety', 'iso 9001'
            ],
            'retail': [
                'retail management', 'merchandising', 'inventory control', 'customer service', 'sales floor',
                'visual merchandising', 'store operations', 'point of sale', 'retail analytics'
            ],
            
            # ZIMBABWE SPECIFIC
            'zimbabwe_skills': [
                'zimswitch', 'ecocash', 'onedollar', 'rtgs', 'forex', 'zimra', 'rbz', 'zimbabwean market',
                'local regulations', 'zimbabwe business environment', 'local banking', 'zse', 'zimbabwe stock exchange'
            ]
        }
        
        # Enhanced qualification levels with comprehensive patterns
        self.qualification_levels = {
            'phd': [
                r'\bph\.?d\.?\b', r'\bdoctorate\b', r'\bdoctor of philosophy\b', r'\bdphil\b',
                r'\bdoctoral\b', r'\bphd candidate\b'
            ],
            'masters': [
                r'\bmaster\b', r'\bmsc?\b', r'\bma\b', r'\bmba\b', r'\bmpa\b', r'\bllm\b', r'\bmed\b',
                r'\bpostgraduate\b', r'\bgraduate degree\b', r'\bmsc\b', r'\bmasters\b', r'\bmaster\'s\b'
            ],
            'bachelors': [
                r'\bbachelor\b', r'\bbsc?\b', r'\bba\b', r'\bbcom\b', r'\bbeng\b', r'\bundergraduate\b',
                r'\bdegree\b', r'\bbachelor\'s\b', r'\bbtech\b', r'\bbachelor of\b'
            ],
            'diploma': [
                r'\bdiploma\b', r'\badvanced diploma\b', r'\bhigher diploma\b', r'\bnational diploma\b',
                r'\bpostgraduate diploma\b'
            ],
            'certificate': [
                r'\bcertificate\b', r'\bprofessional certificate\b', r'\bvocational certificate\b',
                r'\bcertification\b', r'\bcertified\b'
            ],
            'high_school': [
                r'\ba level\b', r'\bo level\b', r'\bhigh school\b', r'\bsecondary education\b',
                r'\badvanced level\b', r'\bordinary level\b', r'\bsecondary school\b'
            ]
        }
        
        # Industry sectors for experience matching
        self.industry_sectors = {
            'technology': ['software', 'it', 'technology', 'computer', 'programming', 'developer', 'engineer', 'tech'],
            'finance': ['banking', 'finance', 'investment', 'accounting', 'audit', 'financial', 'bank', 'investment'],
            'healthcare': ['medical', 'healthcare', 'hospital', 'pharmaceutical', 'nursing', 'doctor', 'health', 'clinic'],
            'education': ['education', 'teaching', 'academic', 'university', 'school', 'lecturer', 'college', 'institution'],
            'manufacturing': ['manufacturing', 'production', 'factory', 'industrial', 'engineering', 'plant', 'assembly'],
            'retail': ['retail', 'sales', 'customer service', 'merchandising', 'store', 'shop', 'outlet'],
            'marketing': ['marketing', 'advertising', 'brand', 'digital marketing', 'social media', 'promotion'],
            'government': ['government', 'public sector', 'municipal', 'civil service', 'public service'],
            'consulting': ['consulting', 'consultant', 'advisory', 'strategy consulting'],
            'nonprofit': ['nonprofit', 'ngo', 'non-governmental', 'charity', 'non profit']
        }

    def generate_confidential_code(self, phone_number: str, text: str) -> str:
        """Generate unique confidential code for resume using phone number"""
        if phone_number:
            # Use phone number hash for confidentiality
            phone_hash = hashlib.sha256(phone_number.encode()).hexdigest()[:8].upper()
            timestamp = datetime.now().strftime("%m%d")
            return f"CV-{phone_hash}-{timestamp}"
        else:
            # Fallback: use text content hash
            text_hash = hashlib.md5(text[:500].encode()).hexdigest()[:8].upper()
            return f"CV-{text_hash}"

    def extract_text_from_file(self, file_path: str, file_type: str) -> str:
        """Extract text from PDF or DOCX files"""
        text = ""
        try:
            if file_type.lower() == "pdf":
                reader = PdfReader(file_path)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            elif file_type.lower() == "docx":
                doc = Document(file_path)
                for para in doc.paragraphs:
                    text += para.text + "\n"
                # Extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                        text += "\n"
        except Exception as e:
            print(f"Error extracting text: {e}")
        return text

    def extract_phone_number(self, text: str) -> str:
        """Extract phone number for confidential coding"""
        phone_patterns = [
            r'(\+263\s?\d{2}\s?\d{3}\s?\d{3,4})',
            r'(07[1-8]\s?\d{3}\s?\d{3,4})',
            r'(\+?\d{1,3}[-.\s]?)?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
            r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',
            r'\b\d{4}[-.\s]?\d{3}[-.\s]?\d{3}\b'
        ]
        
        for pattern in phone_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if isinstance(match, tuple):
                    match = match[0]
                clean_phone = re.sub(r'[^\d+]', '', match)
                if 7 <= len(clean_phone) <= 15:
                    return clean_phone
        return ""

    def _extract_dates(self, text: str) -> str:
        """Extract date ranges from text"""
        date_patterns = [
            r'(20\d{2}[-–]\s*(?:20\d{2}|Present|Current|Now))',
            r'(20\d{2})\s*[-–]\s*(20\d{2}|Present|Current|Now)',
            r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s*\d{4}[-–]\s*(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s*\d{4}|Present|Current|Now',
            r'\b(?:19|20)\d{2}\b'
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group()
        return ""

    def extract_qualifications_advanced(self, text: str) -> Dict:
        """Advanced educational qualifications extraction"""
        qualifications = {
            'highest_qualification': '',
            'qualification_level': '',
            'all_qualifications': [],
            'institutions': [],
            'fields_of_study': [],
            'qualification_details': []
        }
        
        lines = text.split('\n')
        education_section = False
        
        for i, line in enumerate(lines):
            line_clean = line.strip()
            line_lower = line_clean.lower()
            
            # Detect education section with more keywords
            if any(keyword in line_lower for keyword in [
                'education', 'qualifications', 'academic', 'degrees', 
                'educational background', 'academic qualifications'
            ]):
                education_section = True
                continue
            
            if education_section:
                # Check for end of education section
                if any(section in line_lower for section in [
                    'experience', 'work experience', 'employment', 
                    'skills', 'projects', 'certifications', 'professional'
                ]) and len(line_clean) < 50:
                    break
                
                # Extract qualification with enhanced patterns
                qual_info = self._extract_qualification_details(line_clean)
                if qual_info:
                    qualifications['all_qualifications'].append(qual_info)
                    qualifications['qualification_details'].append(qual_info)
        
        # Determine highest qualification
        if qualifications['all_qualifications']:
            level_priority = {'phd': 6, 'masters': 5, 'bachelors': 4, 'diploma': 3, 'certificate': 2, 'high_school': 1}
            highest_qual = max(qualifications['all_qualifications'], 
                             key=lambda x: level_priority.get(x['level'], 0))
            qualifications['highest_qualification'] = highest_qual['text']
            qualifications['qualification_level'] = highest_qual['level']
            qualifications['institutions'] = [q['institution'] for q in qualifications['all_qualifications'] if q['institution']]
            qualifications['fields_of_study'] = [q['field'] for q in qualifications['all_qualifications'] if q['field']]
        
        return qualifications

    def _extract_qualification_details(self, text: str) -> Dict:
        """Extract detailed qualification information"""
        qual_level = self._detect_qualification_level_advanced(text)
        if not qual_level:
            return None
        
        institution = self._extract_institution_advanced(text)
        field = self._extract_field_of_study_advanced(text)
        dates = self._extract_dates(text)
        
        return {
            'text': text,
            'level': qual_level,
            'institution': institution,
            'field': field,
            'dates': dates,
            'confidence': self._calculate_qualification_confidence(text, qual_level)
        }

    def _detect_qualification_level_advanced(self, text: str) -> str:
        """Advanced qualification level detection"""
        text_lower = text.lower()
        
        for level, patterns in self.qualification_levels.items():
            for pattern in patterns:
                if re.search(pattern, text_lower):
                    return level
        return ''

    def _extract_institution_advanced(self, text: str) -> str:
        """Extract educational institution with enhanced patterns"""
        # Common university patterns in Zimbabwe and internationally
        university_patterns = [
            r'University of (\w+\s?){1,3}',
            r'(\w+\s?){1,3}University',
            r'(\w+\s?){1,3}College',
            r'(\w+\s?){1,3}Institute',
            r'(\w+\s?){1,3}School',
            r'(\w+\s?){1,3}Polytechnic',
            r'Harare Institute of Technology',
            r'University of Zimbabwe',
            r'National University of Science and Technology',
            r'Africa University',
            r'Midlands State University',
            r'Great Zimbabwe University',
            r'Chinhoyi University of Technology',
            r'Bindura University',
            r'Lupane State University',
            r'Zimbabwe Open University'
        ]
        
        for pattern in university_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group().strip()
        return ''

    def _extract_field_of_study_advanced(self, text: str) -> str:
        """Extract field of study with comprehensive field list"""
        fields = [
            'computer science', 'information systems', 'business administration', 'engineering',
            'medicine', 'law', 'accounting', 'finance', 'marketing', 'economics', 'mathematics',
            'physics', 'chemistry', 'biology', 'psychology', 'sociology', 'political science',
            'information technology', 'software engineering', 'data science', 'artificial intelligence',
            'mechanical engineering', 'electrical engineering', 'civil engineering', 'chemical engineering',
            'public health', 'nursing', 'pharmacy', 'dentistry', 'education', 'teaching',
            'human resources', 'management', 'entrepreneurship', 'international business',
            'journalism', 'communications', 'public relations', 'graphic design', 'architecture'
        ]
        
        text_lower = text.lower()
        for field in fields:
            if field in text_lower:
                return field.title()
        return ''

    def _calculate_qualification_confidence(self, text: str, level: str) -> float:
        """Calculate confidence score for qualification detection"""
        confidence = 0.6  # Base confidence
        
        # Increase confidence if institution is found
        if self._extract_institution_advanced(text):
            confidence += 0.2
        
        # Increase confidence if field of study is found
        if self._extract_field_of_study_advanced(text):
            confidence += 0.1
        
        # Increase confidence if dates are present
        if self._extract_dates(text):
            confidence += 0.1
        
        return min(confidence, 1.0)

    def extract_skills_comprehensive(self, text: str) -> Dict:
        """Comprehensive skill extraction across all categories"""
        found_skills = {}
        text_lower = text.lower()
        total_skill_occurrences = 0
        
        for category, skills in self.skill_keywords.items():
            category_skills = []
            for skill in skills:
                # Use word boundaries for exact matching
                pattern = r'\b' + re.escape(skill) + r'\b'
                matches = re.findall(pattern, text_lower)
                if matches:
                    confidence = self._calculate_skill_confidence_advanced(skill, text_lower, len(matches))
                    frequency = len(matches)
                    total_skill_occurrences += frequency
                    
                    category_skills.append({
                        'skill': skill,
                        'confidence': confidence,
                        'frequency': frequency,
                        'category': category
                    })
            
            if category_skills:
                found_skills[category] = category_skills
        
        # Calculate overall skill metrics
        total_skills = sum(len(skills) for skills in found_skills.values())
        skill_diversity = len(found_skills)  # Number of categories
        
        return {
            'skills_by_category': found_skills,
            'total_skills': total_skills,
            'skill_diversity': skill_diversity,
            'total_skill_occurrences': total_skill_occurrences,
            'top_skills': self._get_top_skills_advanced(found_skills),
            'skill_categories_present': list(found_skills.keys())
        }

    def _calculate_skill_confidence_advanced(self, skill: str, text: str, frequency: int) -> float:
        """Advanced skill confidence calculation"""
        confidence = 0.5  # Base confidence
        
        # Increase if in skills section
        if re.search(r'(skills|technical skills|competencies).*?' + re.escape(skill), text, re.IGNORECASE | re.DOTALL):
            confidence += 0.3
        
        # Increase if in bullet points or listed format
        if re.search(r'[•\-*]\s*.*' + re.escape(skill), text):
            confidence += 0.1
        
        # Increase based on frequency
        if frequency > 1:
            confidence += min(0.1, frequency * 0.05)
        
        return min(confidence, 1.0)

    def _get_top_skills_advanced(self, skills_by_category: Dict) -> List[Dict]:
        """Get top skills by frequency and confidence"""
        all_skills = []
        for category, skills in skills_by_category.items():
            for skill_data in skills:
                all_skills.append({
                    'skill': skill_data['skill'],
                    'category': skill_data['category'],
                    'score': skill_data['confidence'] * skill_data['frequency'],
                    'frequency': skill_data['frequency'],
                    'confidence': skill_data['confidence']
                })
        
        all_skills.sort(key=lambda x: x['score'], reverse=True)
        return all_skills[:15]  # Return top 15 skills

    def extract_experience_advanced(self, text: str) -> Dict:
        """Advanced experience analysis with detailed career progression"""
        experience = {
            'total_experience_years': 0,
            'experience_level': '',
            'industry_sectors': [],
            'key_achievements': [],
            'career_progression': [],
            'companies_worked': [],
            'positions_held': [],
            'career_gaps': 0,
            'promotion_trajectory': 'stable'
        }
        
        lines = text.split('\n')
        experience_section = False
        current_position = {}
        positions = []
        
        for line in lines:
            line_clean = line.strip()
            line_lower = line_clean.lower()
            
            # Detect experience section
            if any(keyword in line_lower for keyword in [
                'experience', 'work experience', 'employment', 'career', 
                'professional experience', 'work history'
            ]):
                experience_section = True
                continue
            
            if experience_section:
                # Detect end of experience section
                if any(section in line_lower for section in [
                    'education', 'skills', 'projects', 'certifications', 
                    'awards', 'references', 'personal'
                ]) and len(line_clean) < 50:
                    break
                
                # Extract position information
                position_info = self._extract_position_info(line_clean)
                if position_info:
                    if current_position:
                        positions.append(current_position)
                    current_position = position_info
                
                # Extract achievements
                if self._is_achievement_advanced(line_clean):
                    achievement = self._clean_achievement(line_clean)
                    if achievement:
                        experience['key_achievements'].append(achievement)
                
                # Detect industry sectors
                sector = self._detect_industry_sector_advanced(line_clean)
                if sector and sector not in experience['industry_sectors']:
                    experience['industry_sectors'].append(sector)
        
        # Add the last position
        if current_position:
            positions.append(current_position)
        
        experience['career_progression'] = positions
        experience['companies_worked'] = list(set([p.get('company', '') for p in positions if p.get('company')]))
        experience['positions_held'] = list(set([p.get('position', '') for p in positions if p.get('position')]))
        
        # Calculate experience metrics
        experience['total_experience_years'] = self._calculate_total_experience_advanced(positions)
        experience['experience_level'] = self._determine_experience_level_advanced(experience['total_experience_years'])
        experience['career_gaps'] = self._detect_career_gaps(positions)
        experience['promotion_trajectory'] = self._analyze_career_trajectory(positions)
        
        return experience

    def _extract_position_info(self, text: str) -> Dict:
        """Extract detailed position information"""
        dates = self._extract_dates(text)
        if not dates:
            return None
        
        position = self._extract_position_title(text)
        company = self._extract_company_name(text)
        
        return {
            'position': position,
            'company': company,
            'dates': dates,
            'duration_years': self._calculate_position_duration(dates),
            'description': text
        }

    def _extract_position_title(self, text: str) -> str:
        """Extract position title"""
        # Common position patterns
        position_patterns = [
            r'^(.*?)(?=\s*(?:at|,|\d))',
            r'(?:position|role|title)[:\s]+(.*?)(?=\s|$)',
            r'^(.*?)(?=\s*-\s*)'
        ]
        
        for pattern in position_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                title = match.group(1).strip()
                if len(title) > 3 and len(title) < 100:
                    return title
        return ''

    def _extract_company_name(self, text: str) -> str:
        """Extract company name"""
        company_patterns = [
            r'at\s+([^,]+?)(?=\s|$)',
            r',\s+([^,]+?)(?=\s|$)',
            r'-\s+([^,]+?)(?=\s|$)',
            r'company[:\s]+(.*?)(?=\s|$)'
        ]
        
        for pattern in company_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                company = match.group(1).strip()
                if len(company) > 2 and len(company) < 100:
                    return company
        return ''

    def _calculate_position_duration(self, dates: str) -> float:
        """Calculate duration of a position in years"""
        year_matches = re.findall(r'20\d{2}|19\d{2}', dates)
        if len(year_matches) >= 2:
            start_year = int(year_matches[0])
            end_year = int(year_matches[1]) if year_matches[1].isdigit() else datetime.now().year
            return end_year - start_year
        return 0.0

    def _is_achievement_advanced(self, text: str) -> bool:
        """Check if line describes an achievement"""
        achievement_indicators = [
            r'achieved', r'increased', r'reduced', r'improved', r'led', r'managed', r'developed',
            r'implemented', r'saved', r'won', r'awarded', r'recognized', r'spearheaded',
            r'created', r'built', r'launched', r'optimized', r'streamlined', r'enhanced',
            r'delivered', r'exceeded', r'accomplished', r'completed', r'successfully'
        ]
        
        text_lower = text.lower()
        return any(re.search(indicator, text_lower) for indicator in achievement_indicators)

    def _clean_achievement(self, text: str) -> str:
        """Clean and format achievement text"""
        # Remove bullet points and extra whitespace
        cleaned = re.sub(r'^[•\-*\s]+', '', text.strip())
        return cleaned if len(cleaned) > 10 else ''  # Filter out very short achievements

    def _detect_industry_sector_advanced(self, text: str) -> str:
        """Advanced industry sector detection"""
        text_lower = text.lower()
        for sector, keywords in self.industry_sectors.items():
            if any(keyword in text_lower for keyword in keywords):
                return sector
        return ''

    def _calculate_total_experience_advanced(self, career_progression: List[Dict]) -> int:
        """Calculate total years of experience with overlap handling"""
        if not career_progression:
            return 0
        
        # Simple calculation for demo - in production, use proper date range merging
        total_years = 0
        for position in career_progression:
            duration = position.get('duration_years', 0)
            if duration > 0:
                total_years += duration
        
        return min(total_years, 50)  # Cap at 50 years

    def _determine_experience_level_advanced(self, years: int) -> str:
        """Determine experience level based on years"""
        if years < 1:
            return "Entry Level (0-1 years)"
        elif years < 3:
            return "Junior Level (1-3 years)"
        elif years < 7:
            return "Mid Level (3-7 years)"
        elif years < 15:
            return "Senior Level (7-15 years)"
        else:
            return "Executive Level (15+ years)"

    def _detect_career_gaps(self, positions: List[Dict]) -> int:
        """Detect career gaps (simplified)"""
        if len(positions) < 2:
            return 0
        
        # Simple gap detection - count positions with very short durations
        short_positions = sum(1 for pos in positions if pos.get('duration_years', 0) < 0.5)
        return short_positions

    def _analyze_career_trajectory(self, positions: List[Dict]) -> str:
        """Analyze career progression trajectory"""
        if len(positions) < 2:
            return "Early Career"
        
        position_count = len(positions)
        avg_duration = sum(pos.get('duration_years', 0) for pos in positions) / position_count
        
        if avg_duration >= 3.0:
            return "Stable Progression"
        elif avg_duration >= 1.5:
            return "Steady Growth"
        else:
            return "Rapid Movement"

    def extract_awards_certifications_advanced(self, text: str) -> Dict:
        """Extract awards, certifications, and honors with details"""
        awards_certifications = {
            'awards': [],
            'certifications': [],
            'honors': [],
            'publications': [],
            'patents': []
        }
        
        lines = text.split('\n')
        
        award_indicators = {
            'awards': ['award', 'prize', 'recognition', 'achievement award'],
            'certifications': ['certification', 'certified', 'license', 'accreditation'],
            'honors': ['honor', 'scholarship', 'fellowship', 'dean\'s list'],
            'publications': ['publication', 'paper', 'journal', 'conference'],
            'patents': ['patent', 'invention', 'intellectual property']
        }
        
        for line in lines:
            line_clean = line.strip()
            line_lower = line_clean.lower()
            
            for category, indicators in award_indicators.items():
                if any(indicator in line_lower for indicator in indicators):
                    awards_certifications[category].append({
                        'text': line_clean,
                        'confidence': 0.8,
                        'type': category
                    })
        
        return awards_certifications

    def parse_resume(self, file_path: str, file_type: str) -> Dict:
        """Main method to parse and analyze resume with enhanced features"""
        try:
            # Extract text
            raw_text = self.extract_text_from_file(file_path, file_type)
            if not raw_text:
                return {"error": "Could not extract text from file"}
            
            # Extract phone number for confidential coding
            phone_number = self.extract_phone_number(raw_text)
            
            # Generate confidential code
            resume_code = self.generate_confidential_code(phone_number, raw_text)
            
            # Extract all information with enhanced methods
            qualifications = self.extract_qualifications_advanced(raw_text)
            skills = self.extract_skills_comprehensive(raw_text)
            experience = self.extract_experience_advanced(raw_text)
            awards_certifications = self.extract_awards_certifications_advanced(raw_text)
            
            # Generate skills list for database storage
            skills_list = []
            for category, skill_data in skills['skills_by_category'].items():
                for skill_info in skill_data:
                    skills_list.append((
                        skill_info['skill'],
                        skill_info['category'],
                        skill_info['confidence']
                    ))
            
            return {
                "resume_code": resume_code,
                "phone_number": phone_number,  # For confidential lookup
                "qualifications": qualifications,
                "skills": skills,
                "experience": experience,
                "awards_certifications": awards_certifications,
                "skills_list": skills_list,
                "summary": self.generate_comprehensive_summary(qualifications, skills, experience, awards_certifications),
                "processed_text": raw_text[:1500],  # Store limited text for reference
                "extraction_timestamp": datetime.now().isoformat(),
                "analysis_metadata": {
                    "skill_categories_found": len(skills['skills_by_category']),
                    "total_skills_identified": skills['total_skills'],
                    "experience_years": experience['total_experience_years'],
                    "qualification_level": qualifications['qualification_level'],
                    "industries_identified": experience['industry_sectors']
                }
            }
            
        except Exception as e:
            return {"error": f"Error parsing resume: {str(e)}"}

    def generate_comprehensive_summary(self, qualifications: Dict, skills: Dict, experience: Dict, awards: Dict) -> str:
        """Generate professional summary focusing on key strengths"""
        summary_parts = []
        
        # Qualification summary
        if qualifications['qualification_level']:
            level_display = qualifications['qualification_level'].replace('_', ' ').title()
            summary_parts.append(f"{level_display} qualified professional")
        
        # Experience summary
        if experience['total_experience_years'] > 0:
            summary_parts.append(f"with {experience['total_experience_years']} years of {experience['experience_level'].lower()} experience")
        
        # Key skills summary
        top_skills = skills.get('top_skills', [])[:5]
        if top_skills:
            skill_names = [skill['skill'] for skill in top_skills]
            summary_parts.append(f"specializing in {', '.join(skill_names)}")
        
        # Industry experience
        if experience['industry_sectors']:
            industries = ', '.join(experience['industry_sectors'][:3])
            summary_parts.append(f"across {industries} sectors")
        
        # Achievements highlight
        if awards.get('awards') or awards.get('certifications'):
            summary_parts.append("with notable achievements and certifications")
        
        return ' '.join(summary_parts) + "."

# Create alias for backward compatibility
ResumeParser = AdvancedResumeAnalyzer